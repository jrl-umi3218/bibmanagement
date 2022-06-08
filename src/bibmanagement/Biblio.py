from bibmanagement.Entry import Entry
from bibmanagement.Entry import Proceedings
from bibmanagement.bibParser import openBib
from bibmanagement.formatExpr import Formatter
from bibmanagement.utils import FormattedString
from bibmanagement.utils import Sort as us
from openpyxl import Workbook
from openpyxl import utils
from openpyxl.styles import Font
from docx import Document
from docx.shared import Pt
import copy

class Biblio:
    def __init__(self, entryList):
        self.entries = []
        self._dictMap = {}
        for e in entryList:
            self.addEntry(e)

    @classmethod
    def fromdb(cls, db):
        bib = cls([])
        for e in db.entries:
            try:
                bib.addEntry(e)
            except ValueError as err:
                print("ValueError: {0}".format(err))
                print(e)
            except BaseException as err:
                print("BaseException: {0}".format(err))
                print(e)
            except:
                print("Unknown error for")
                print(e)
        return bib

    @classmethod
    def fromBibFile(cls, filepath):
        db = openBib(filepath)
        return cls.fromdb(db)
		
    def addEntry(self, e):
        self._dictMap[e['ID']] = len(self.entries)
        self.entries.append(Entry(e))
		
    def __getitem__(self, keyOrIndex):
        if isinstance(keyOrIndex, str):
            return self.entries[self._dictMap[keyOrIndex]]
        else:
            return self.entries[keyOrIndex]
        
    def __contains__(self, key):
        return key in self._dictMap
        
    def __len__(self):
        return len(self.entries)

    def resolveCrossref(self, removeProceedings=True):
        needResolve = False
        for e in self.entries:
            if 'crossref' in e:
                needResolve = True
                break
            
        if needResolve or removeProceedings:
            bib = copy.deepcopy(self)
            for e in bib.entries:
                while 'crossref' in e:
                    c = str(e['crossref'].format())
                    if c not in bib:
                        raise KeyError('Crossref ' + c + ' not found in ' + str(e))
                    del e['crossref']
                    e.integrateCrossref(bib[c])
                    
            if removeProceedings:
                return bib.filter(lambda x: not isinstance(x,Proceedings))
            return bib

    def filter(self, cond):
        '''Return a Biblio with only the entries e for which cond(e) is True'''
        return Biblio([e.raw() for e in self.entries if cond(e)])

    def sort(self, crit):
        '''
        Sort the entries of the bibliography according to the order specify by crit
        
        crit can be:
         - a single format string (as used for Entry.format)
         - a tuple (formatString, order) where order is either \'<\' (ascending) or \'>\' descending
         - a list [c1,...,ck] of the two possibilities above, specifying that the entries are
           sorted first according to c1, then among c1-equal elements according to c2, ...
        '''
        if isinstance(crit, str):
            ff = [Formatter.FixedFormatter(crit)]
            cmp = us.MixComp(['<'])
        elif isinstance(crit,tuple):
            ff = [Formatter.FixedFormatter(crit[0])]
            cmp = us.MixComp([crit[1]])
        else:
            ff = [Formatter.FixedFormatter(c[0]) if isinstance(c,tuple) else Formatter.FixedFormatter(c) for c in crit]
            cmp = us.MixComp([c[1] if isinstance(c,tuple) is tuple else '<' for c in crit])
        self.entries.sort(key = lambda e: us.MixCompKey(tuple(str(f.run(e)) for f in ff), cmp))


    def toExcel(self, format, startRow=1, out='demo'):
        '''
        Output the biblio to an excel file according to a given format
        
        One entry of the biblio will appear on each row, and columns of this
        row will depend on the specified format.
        
        Parameters
        ----------
        format : dict
            A dictionary where each key is the letter(s) of a column, and each
            item is a formatting string specifying the content of that column.
        startRow : int
            Number of the first row in which an entry is written.
        out : str
            The path of the output file, without the extension.
        '''
        wb = Workbook()
        # grab the active worksheet
        ws = wb.active

        for i in range(0,len(self.entries)):
            for col,f in format.items():
                r = self.entries[i].format(f)
                cell = ws.cell(row=startRow+i, column=utils.cell.column_index_from_string(col))
                cell.value=str(r)
                if r.style:
                    s = r.style
                    u = 'single' if s.underline else 'none'
                    cell.font = Font(name=s.font, size=s.fontSize, bold=s.bold, italic=s.italic, underline=u, strikethrough=s.strikethrough)

        # Save the file
        wb.save(out + '.xlsx')

    def toWord(self, format, styleFilters=[], out='demo'):
        document = Document()
        for e in self.entries:
            p = document.add_paragraph(style='List Number 2')
            r = e.format(format)
            for sf in styleFilters:
                r = sf(r)
            for ri in r.singles:
                run = p.add_run(str(ri))
                s = ri.style if ri.style else FormattedString.Style()
                run.bold = s.bold
                run.italic = s.italic
                run.underline = s.underline
                run.font.strike = s.strikethrough
                run.font.name  = s.font
                run.font.size = Pt(s.fontSize)

        document.save(out+'.docx')